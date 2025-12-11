#!/usr/bin/env python3
"""
Convert SETUP_GUIDE.md to a Word document with proper formatting
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import re

def add_markdown_to_docx(md_file, docx_file):
    """Convert Markdown file to Word document"""
    
    # Create a new Document
    doc = Document()
    
    # Set up styles
    styles = doc.styles
    
    # Read the markdown file
    with open(md_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        
        # Skip empty lines
        if not line:
            i += 1
            continue
        
        # Heading 1 (# Title)
        if line.startswith('# '):
            p = doc.add_heading(line[2:], level=1)
            p.runs[0].font.color.rgb = RGBColor(0, 51, 102)
        
        # Heading 2 (## Title)
        elif line.startswith('## '):
            p = doc.add_heading(line[3:], level=2)
            p.runs[0].font.color.rgb = RGBColor(0, 102, 204)
        
        # Heading 3 (### Title)
        elif line.startswith('### '):
            p = doc.add_heading(line[4:], level=3)
            p.runs[0].font.color.rgb = RGBColor(51, 102, 153)
        
        # Horizontal rule (---)
        elif line.startswith('---'):
            doc.add_paragraph()  # Add spacing
        
        # Blockquote (> text)
        elif line.startswith('> '):
            # Handle special blockquotes
            if '[!IMPORTANT]' in line:
                p = doc.add_paragraph()
                run = p.add_run('⚠️ IMPORTANT: ')
                run.bold = True
                run.font.color.rgb = RGBColor(255, 140, 0)
                # Get the rest of the text
                text = line.replace('> [!IMPORTANT]', '').strip()
                if text:
                    p.add_run(text)
            elif '[!CAUTION]' in line:
                p = doc.add_paragraph()
                run = p.add_run('⚠️ CAUTION: ')
                run.bold = True
                run.font.color.rgb = RGBColor(255, 0, 0)
                text = line.replace('> [!CAUTION]', '').strip()
                if text:
                    p.add_run(text)
            else:
                p = doc.add_paragraph(line[2:])
                p.paragraph_format.left_indent = Inches(0.5)
                p.style = 'Intense Quote'
        
        # Unordered list (- item or * item)
        elif line.startswith('- ') or line.startswith('* '):
            text = line[2:]
            # Handle checkboxes
            if text.startswith('[ ] '):
                text = '☐ ' + text[4:]
            elif text.startswith('[x] '):
                text = '☑ ' + text[4:]
            
            p = doc.add_paragraph(text, style='List Bullet')
            # Apply formatting to inline code and bold
            format_inline_styles(p)
        
        # Ordered list (1. item)
        elif re.match(r'^\d+\.\s', line):
            text = re.sub(r'^\d+\.\s', '', line)
            p = doc.add_paragraph(text, style='List Number')
            format_inline_styles(p)
        
        # Code block (```)
        elif line.startswith('```'):
            # Collect all lines until closing ```
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i].rstrip())
                i += 1
            
            # Add code block
            if code_lines:
                code_text = '\n'.join(code_lines)
                p = doc.add_paragraph(code_text)
                p.style = 'No Spacing'
                for run in p.runs:
                    run.font.name = 'Courier New'
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(0, 0, 0)
                # Add light gray background effect
                p.paragraph_format.left_indent = Inches(0.3)
                p.paragraph_format.right_indent = Inches(0.3)
        
        # Table (| header |)
        elif line.startswith('|'):
            # Collect table rows
            table_rows = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                row = [cell.strip() for cell in lines[i].strip().split('|')[1:-1]]
                # Skip separator rows (|---|---|)
                if not all(re.match(r'^-+$', cell.strip()) for cell in row):
                    table_rows.append(row)
                i += 1
            i -= 1  # Back up one line
            
            if table_rows:
                # Create table
                table = doc.add_table(rows=len(table_rows), cols=len(table_rows[0]))
                table.style = 'Light Grid Accent 1'
                
                # Fill table
                for row_idx, row_data in enumerate(table_rows):
                    for col_idx, cell_data in enumerate(row_data):
                        cell = table.rows[row_idx].cells[col_idx]
                        cell.text = cell_data
                        # Make header row bold
                        if row_idx == 0:
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.bold = True
        
        # Regular paragraph
        else:
            p = doc.add_paragraph(line)
            format_inline_styles(p)
        
        i += 1
    
    # Save the document
    doc.save(docx_file)
    print(f"✅ Successfully converted {md_file} to {docx_file}")

def format_inline_styles(paragraph):
    """Format inline styles like bold, code, etc."""
    text = paragraph.text
    paragraph.clear()
    
    # Pattern for inline code (`code`), bold (**text**), and links
    parts = re.split(r'(`[^`]+`|\*\*[^*]+\*\*|\[[^\]]+\]\([^\)]+\))', text)
    
    for part in parts:
        if not part:
            continue
        
        # Inline code
        if part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(199, 37, 78)
        
        # Bold
        elif part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        
        # Links [text](url)
        elif part.startswith('[') and '](' in part:
            match = re.match(r'\[([^\]]+)\]\(([^\)]+)\)', part)
            if match:
                link_text = match.group(1)
                link_url = match.group(2)
                run = paragraph.add_run(link_text)
                run.font.color.rgb = RGBColor(0, 0, 255)
                run.underline = True
        
        # Regular text
        else:
            paragraph.add_run(part)

if __name__ == '__main__':
    import sys
    
    # Accept command-line arguments or use defaults
    if len(sys.argv) >= 2:
        md_file = sys.argv[1]
        # Generate output filename from input filename
        if len(sys.argv) >= 3:
            docx_file = sys.argv[2]
        else:
            docx_file = md_file.replace('.md', '.docx')
    else:
        md_file = 'SETUP_GUIDE.md'
        docx_file = 'SETUP_GUIDE.docx'
    
    add_markdown_to_docx(md_file, docx_file)
